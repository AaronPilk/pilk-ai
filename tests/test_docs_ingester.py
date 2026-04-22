"""Docs-folder ingester — walk + render + home-scope clamp.

Every test points the scanner at tmp_path and uses tmp_path as the
fake home root so the home-scope check is satisfied without touching
real user data.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from core.integrations.ingesters.docs import (
    DEFAULT_EXTENSIONS,
    DocsIngestError,
    render_doc_note,
    scan_docs,
)


def _mk(root: Path, rel: str, body: str) -> Path:
    path = root / rel
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(body, encoding="utf-8")
    return path


# ── scan_docs ────────────────────────────────────────────────────────


def test_scan_picks_up_text_natives(tmp_path: Path) -> None:
    _mk(tmp_path, "a.md", "# hello")
    _mk(tmp_path, "notes/b.txt", "plain text")
    _mk(tmp_path, "nested/deep/c.json", '{"k": 1}')
    # Non-matching extensions silently skipped (no entry in skipped).
    _mk(tmp_path, "image.jpg", "not-actually-an-image")
    scan = scan_docs(tmp_path, home=tmp_path)
    names = sorted(d.abs_path.name for d in scan.found)
    assert names == ["a.md", "b.txt", "c.json"]


def test_scan_honours_extensions_filter(tmp_path: Path) -> None:
    _mk(tmp_path, "a.md", "md")
    _mk(tmp_path, "b.txt", "txt")
    scan = scan_docs(tmp_path, extensions=(".md",), home=tmp_path)
    names = [d.abs_path.name for d in scan.found]
    assert names == ["a.md"]


def test_scan_respects_max_files(tmp_path: Path) -> None:
    for i in range(10):
        _mk(tmp_path, f"file_{i}.txt", f"body {i}")
    scan = scan_docs(tmp_path, max_files=3, home=tmp_path)
    assert len(scan.found) == 3
    # Overflow documented in skipped.
    assert any("max_files cap" in reason for _, reason in scan.skipped)


def test_scan_skips_hidden_dirs_and_files(tmp_path: Path) -> None:
    _mk(tmp_path, ".git/config", "secret")
    _mk(tmp_path, "legit.md", "ok")
    _mk(tmp_path, ".hidden.txt", "shh")
    scan = scan_docs(tmp_path, home=tmp_path)
    kept = [d.abs_path.name for d in scan.found]
    assert kept == ["legit.md"]


def test_scan_skips_oversized_files(tmp_path: Path) -> None:
    _mk(tmp_path, "small.md", "tiny")
    _mk(tmp_path, "big.md", "x" * (11 * 1024 * 1024))  # 11 MiB, over cap
    scan = scan_docs(tmp_path, home=tmp_path)
    names = [d.abs_path.name for d in scan.found]
    assert names == ["small.md"]
    assert any("too large" in r for _, r in scan.skipped)


def test_scan_non_recursive_stays_at_top_level(tmp_path: Path) -> None:
    _mk(tmp_path, "top.md", "x")
    _mk(tmp_path, "sub/deep.md", "x")
    scan = scan_docs(tmp_path, recursive=False, home=tmp_path)
    names = [d.abs_path.name for d in scan.found]
    assert names == ["top.md"]


def test_scan_rejects_missing_root(tmp_path: Path) -> None:
    with pytest.raises(DocsIngestError, match="not found"):
        scan_docs(tmp_path / "does-not-exist", home=tmp_path)


def test_scan_rejects_file_as_root(tmp_path: Path) -> None:
    f = _mk(tmp_path, "a.md", "x")
    with pytest.raises(DocsIngestError, match="not a directory"):
        scan_docs(f, home=tmp_path)


def test_scan_refuses_path_outside_home(tmp_path: Path) -> None:
    # Fake home is a subdir; passing the parent should get rejected.
    fake_home = tmp_path / "home"
    fake_home.mkdir()
    _mk(fake_home, "ok.md", "x")
    _mk(tmp_path, "outside.md", "x")
    with pytest.raises(DocsIngestError, match="must live under"):
        scan_docs(tmp_path, home=fake_home)


def test_scan_handles_undecodable_bytes(tmp_path: Path) -> None:
    # Pure binary that fails utf-8 AND latin-1 is vanishingly rare —
    # latin-1 decodes anything. We instead exercise the happy
    # fallback: latin-1 decodes a file that isn't valid utf-8.
    path = tmp_path / "weird.txt"
    path.write_bytes(b"\xff\xfe hello \xee\xdd")
    scan = scan_docs(tmp_path, home=tmp_path)
    assert len(scan.found) == 1
    # Text is non-empty — latin-1 succeeded.
    assert scan.found[0].text


def test_scan_uses_default_extension_set(tmp_path: Path) -> None:
    # Plain-text extensions round-trip through the default config.
    # PDF / docx need a real binary structure, so they're covered by
    # the dedicated pypdf / python-docx tests below.
    binary_exts = {".pdf", ".docx"}
    text_exts = [e for e in DEFAULT_EXTENSIONS if e not in binary_exts]
    for suffix in text_exts:
        _mk(tmp_path, f"file{suffix}", "body")
    scan = scan_docs(tmp_path, home=tmp_path)
    assert len(scan.found) == len(text_exts)


# ── render_doc_note ─────────────────────────────────────────────────


def test_render_preserves_markdown_verbatim(tmp_path: Path) -> None:
    _mk(tmp_path, "note.md", "# Topic\n\n- point one\n- [[Wiki Link]]\n")
    scan = scan_docs(tmp_path, home=tmp_path)
    note = render_doc_note(scan.found[0], scan_root=scan.root)
    assert note.title == "note"
    assert note.path.startswith("ingested/docs/")
    assert note.path.endswith(".md")
    assert "# Topic" in note.body
    # Wiki links survive into the vault untouched.
    assert "[[Wiki Link]]" in note.body
    # Frontmatter present.
    assert note.body.startswith("---\n")
    assert "source:" in note.body
    assert "tags: [ingested, docs]" in note.body


def test_render_fences_non_markdown_content(tmp_path: Path) -> None:
    _mk(tmp_path, "data.json", '{"x": 1}')
    scan = scan_docs(tmp_path, home=tmp_path)
    note = render_doc_note(scan.found[0], scan_root=scan.root)
    assert "```json" in note.body
    assert '{"x": 1}' in note.body


def test_render_strips_html_tags(tmp_path: Path) -> None:
    _mk(
        tmp_path,
        "page.html",
        "<html><body><h1>Hi</h1><p>para &amp; stuff</p></body></html>",
    )
    scan = scan_docs(tmp_path, home=tmp_path)
    note = render_doc_note(scan.found[0], scan_root=scan.root)
    # Tags removed; entities decoded.
    assert "<h1>" not in note.body
    assert "Hi" in note.body
    assert "para & stuff" in note.body


def test_render_preserves_folder_layout_in_vault_path(tmp_path: Path) -> None:
    _mk(tmp_path, "Projects/Skyway/notes.md", "x")
    scan = scan_docs(tmp_path, home=tmp_path)
    note = render_doc_note(scan.found[0], scan_root=scan.root)
    # Slugged path mirrors the source tree so Obsidian graph
    # clusters hold together.
    assert note.path == "ingested/docs/projects/skyway/notes.md"


def test_render_slugs_each_path_segment(tmp_path: Path) -> None:
    _mk(tmp_path, "Client Files/My Notes.md", "x")
    scan = scan_docs(tmp_path, home=tmp_path)
    note = render_doc_note(scan.found[0], scan_root=scan.root)
    # Spaces become dashes; casing folds; suffix preserved.
    assert note.path == "ingested/docs/client-files/my-notes.md"


# ── IGNORED_DIR_NAMES + wide walk ────────────────────────────────────


def test_scan_prunes_known_cache_folders(tmp_path: Path) -> None:
    # The walker must skip Library / node_modules / .git etc. even
    # when they contain text files — otherwise a home walk drowns
    # in machine-generated junk.
    _mk(tmp_path, "project/src.md", "kept")
    _mk(tmp_path, "project/node_modules/react/package.json", "{}")
    _mk(tmp_path, "Library/Caches/whatever/log.txt", "noise")
    _mk(tmp_path, ".git/config", "secret")
    _mk(tmp_path, "dist/bundle.json", "{}")
    scan = scan_docs(tmp_path, home=tmp_path)
    kept_paths = {str(d.abs_path) for d in scan.found}
    # The one legit file lands.
    assert any(p.endswith("src.md") for p in kept_paths)
    # Everything under a pruned dir is absent.
    for bad in ("node_modules", "Library", ".git", "dist"):
        assert not any(bad in p for p in kept_paths), bad


def test_scan_ignores_pilk_brain_itself(tmp_path: Path) -> None:
    # The operator's vault name. Walking it recursively would copy
    # the vault into the vault, over and over.
    _mk(tmp_path, "user-note.md", "kept")
    _mk(tmp_path, "PILK-brain/ingested/docs/already.md", "loop")
    scan = scan_docs(tmp_path, home=tmp_path)
    kept_paths = {str(d.abs_path) for d in scan.found}
    assert any(p.endswith("user-note.md") for p in kept_paths)
    assert not any("PILK-brain" in p for p in kept_paths)


# ── PDF + docx support ──────────────────────────────────────────────


def test_scan_picks_up_pdf_when_pypdf_available(tmp_path: Path) -> None:
    # pypdf transitively imports `cryptography`, which can blow up via
    # a Rust-level panic on some sandboxed interpreters (missing
    # `_cffi_backend`). Catch BaseException so the panic counts as a
    # skip, not a failure.
    try:
        from pypdf import PdfWriter
    except BaseException as e:
        pytest.skip(f"pypdf unavailable in this env: {e}")

    try:
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        pdf_path = tmp_path / "invoice.pdf"
        with pdf_path.open("wb") as fh:
            writer.write(fh)
    except BaseException as e:
        pytest.skip(f"pypdf write failed in this env: {e}")
    scan = scan_docs(tmp_path, home=tmp_path)
    paths = [d.abs_path.name for d in scan.found]
    assert "invoice.pdf" in paths


def test_scan_picks_up_docx_when_python_docx_available(
    tmp_path: Path,
) -> None:
    pytest.importorskip("docx")
    import docx

    docx_path = tmp_path / "proposal.docx"
    doc = docx.Document()
    doc.add_paragraph("First paragraph of the proposal.")
    doc.add_paragraph("Second paragraph with more content.")
    doc.save(str(docx_path))

    scan = scan_docs(tmp_path, home=tmp_path)
    found = [d for d in scan.found if d.abs_path.name == "proposal.docx"]
    assert found, "docx wasn't picked up"
    body = found[0].text
    assert "First paragraph" in body
    assert "Second paragraph" in body


# ── cap bump ────────────────────────────────────────────────────────


def test_hard_max_files_supports_full_home_walk() -> None:
    # Sanity-check the constant so nobody silently lowers it back to
    # 10K and re-breaks the "ingest my whole home" use case.
    from core.integrations.ingesters.docs import HARD_MAX_FILES

    assert HARD_MAX_FILES >= 50_000
