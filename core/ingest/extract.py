"""Text extraction for the ingestion pipeline.

Supported formats (and whatever fallback applies):

  - ``.txt`` / ``.md``     — read as utf-8 (latin-1 fallback)
  - ``.pdf``                 — pypdf
  - ``.docx``                — python-docx
  - ``.csv`` / ``.tsv``      — read as text; delimiter preserved
  - ``.xlsx``                — openpyxl when available; else
                               ``ExtractionError`` with a clear hint
  - ``.html``/``.htm``       — strip tags via stdlib (best-effort)
  - ``.rtf``                 — best-effort plain-text fallback
  - everything else          — ``ExtractionError``

Image / audio / video extraction is deferred — those need OCR /
transcription pipelines that aren't in scope here.
"""

from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path

from core.logging import get_logger

log = get_logger("pilkd.ingest.extract")


_TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".log"}
_TABULAR_SUFFIXES = {".csv", ".tsv"}
_HTML_SUFFIXES = {".html", ".htm"}


class ExtractionError(RuntimeError):
    """Raised when a file cannot be extracted to text. The pipeline
    catches this, marks the row failed with the error message, and
    moves the file to the failed/ archive."""


@dataclass
class ExtractedText:
    text: str
    file_type: str  # short label: txt|md|pdf|docx|csv|xlsx|html|rtf|other
    pages: int | None = None  # for PDFs
    metadata: dict | None = None


def supported_extensions() -> list[str]:
    return sorted(
        list(_TEXT_SUFFIXES)
        + list(_TABULAR_SUFFIXES)
        + list(_HTML_SUFFIXES)
        + [".pdf", ".docx", ".xlsx", ".rtf"]
    )


def extract_text(path: Path) -> ExtractedText:
    """Extract text from ``path`` based on its extension.

    Pure synchronous — embedding / writing happens in the pipeline.
    """
    suf = path.suffix.lower()
    if suf in _TEXT_SUFFIXES:
        return ExtractedText(
            text=_read_text_file(path),
            file_type=("md" if suf in (".md", ".markdown") else "txt"),
        )
    if suf == ".pdf":
        return _read_pdf(path)
    if suf == ".docx":
        return _read_docx(path)
    if suf in _TABULAR_SUFFIXES:
        return _read_csv(path, suffix=suf)
    if suf == ".xlsx":
        return _read_xlsx(path)
    if suf in _HTML_SUFFIXES:
        return _read_html(path)
    if suf == ".rtf":
        return _read_rtf(path)
    raise ExtractionError(
        f"unsupported extension {suf!r}; supported: "
        f"{', '.join(supported_extensions())}"
    )


# ── Implementations ──────────────────────────────────────────────


def _read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ExtractionError(f"cannot decode {path.name} as text")


def _read_pdf(path: Path) -> ExtractedText:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ExtractionError(
            "pypdf not installed; cannot extract PDF text"
        ) from e
    try:
        reader = PdfReader(str(path))
    except Exception as e:
        raise ExtractionError(f"could not open PDF: {e}") from e
    parts: list[str] = []
    page_count = 0
    for page in reader.pages:
        page_count += 1
        try:
            t = page.extract_text() or ""
        except Exception:
            continue
        if t.strip():
            parts.append(t)
    text = "\n\n".join(parts)
    if not text.strip():
        # Encrypted / scanned / image-only PDFs return empty. We
        # surface a stub so the operator knows to OCR or re-export.
        raise ExtractionError(
            f"PDF appears to be image-only or encrypted "
            f"({page_count} pages, no extractable text)"
        )
    return ExtractedText(
        text=text, file_type="pdf", pages=page_count,
    )


def _read_docx(path: Path) -> ExtractedText:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise ExtractionError(
            "python-docx not installed; cannot extract DOCX text"
        ) from e
    try:
        doc = docx.Document(str(path))
    except Exception as e:
        raise ExtractionError(f"could not open DOCX: {e}") from e
    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts)
    if not text:
        raise ExtractionError("DOCX is empty")
    return ExtractedText(text=text, file_type="docx")


def _read_csv(path: Path, *, suffix: str) -> ExtractedText:
    delim = "\t" if suffix == ".tsv" else ","
    raw = _read_text_file(path)
    # Preserve as plain text — markdown-friendly because tables
    # already render in Obsidian. We don't try to "interpret" the
    # data here; downstream summarization can do that.
    rows: list[list[str]] = []
    reader = csv.reader(io.StringIO(raw), delimiter=delim)
    for r in reader:
        rows.append(r)
    if not rows:
        raise ExtractionError("CSV has no rows")
    # Render as a markdown-style table for the first 100 rows so
    # the resulting brain note is actually readable. Beyond 100
    # rows the raw CSV is included as a code block.
    out: list[str] = []
    if len(rows) <= 100:
        header = rows[0]
        out.append("| " + " | ".join(header) + " |")
        out.append("|" + "|".join(["---"] * len(header)) + "|")
        for r in rows[1:]:
            out.append(
                "| "
                + " | ".join((c or "").replace("|", "\\|") for c in r)
                + " |"
            )
    else:
        out.append(
            f"```{('tsv' if suffix == '.tsv' else 'csv')}\n{raw}\n```"
        )
    return ExtractedText(
        text="\n".join(out),
        file_type=("tsv" if suffix == ".tsv" else "csv"),
        metadata={"rows": len(rows)},
    )


def _read_xlsx(path: Path) -> ExtractedText:
    try:
        import openpyxl
    except ImportError as e:
        raise ExtractionError(
            "openpyxl not installed; cannot extract XLSX text. "
            "Install with: uv add openpyxl"
        ) from e
    try:
        wb = openpyxl.load_workbook(str(path), data_only=True)
    except Exception as e:
        raise ExtractionError(f"could not open XLSX: {e}") from e
    parts: list[str] = []
    sheet_count = 0
    for sheet in wb.worksheets:
        sheet_count += 1
        parts.append(f"## {sheet.title}")
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        if len(rows) <= 100:
            header = [str(c or "") for c in rows[0]]
            parts.append("| " + " | ".join(header) + " |")
            parts.append("|" + "|".join(["---"] * len(header)) + "|")
            for r in rows[1:]:
                cells = [str(c) if c is not None else "" for c in r]
                parts.append(
                    "| "
                    + " | ".join(c.replace("|", "\\|") for c in cells)
                    + " |"
                )
        else:
            parts.append(f"_(sheet has {len(rows)} rows; truncated)_")
            for r in rows[:100]:
                parts.append(
                    "  ".join(
                        str(c) if c is not None else "" for c in r
                    )
                )
    return ExtractedText(
        text="\n".join(parts),
        file_type="xlsx",
        metadata={"sheets": sheet_count},
    )


class _HTMLToText(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []
        self._skip = 0

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style", "noscript"):
            self._skip += 1

    def handle_endtag(self, tag):
        if tag in ("script", "style", "noscript") and self._skip:
            self._skip -= 1

    def handle_data(self, data):
        if self._skip:
            return
        text = data.strip()
        if text:
            self._parts.append(text)

    def text(self) -> str:
        return "\n".join(self._parts)


def _read_html(path: Path) -> ExtractedText:
    raw = _read_text_file(path)
    parser = _HTMLToText()
    parser.feed(raw)
    out = parser.text()
    if not out:
        raise ExtractionError("HTML produced no extractable text")
    return ExtractedText(text=out, file_type="html")


def _read_rtf(path: Path) -> ExtractedText:
    """Best-effort: strip RTF control words. Good enough for plain
    notes; complex rtf with tables loses formatting. Real-world
    operators usually export to docx anyway."""
    raw = _read_text_file(path)
    # Drop control words like \word and groups that don't carry text.
    stripped = re.sub(r"\\[a-zA-Z]+\d* ?", "", raw)
    stripped = re.sub(r"[{}]", "", stripped)
    stripped = re.sub(r"\\[\'\"]", "", stripped)
    cleaned = "\n".join(
        ln.strip() for ln in stripped.splitlines() if ln.strip()
    )
    if not cleaned:
        raise ExtractionError("RTF produced no text")
    return ExtractedText(text=cleaned, file_type="rtf")


__all__ = [
    "ExtractedText",
    "ExtractionError",
    "extract_text",
    "supported_extensions",
]
